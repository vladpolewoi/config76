#!/usr/bin/env python3
"""Render a DR.md (Заключение разработки) into a .docx matching the manager's
house style:
  - title 18pt bold
  - Раздел headings: Heading 2, bold, blue 2F75B5
  - meta table (first 2-col table): grey D9D9D9 bold label column
  - Раздел A/B tables: dark-blue 1F4E79 header row, white bold text
  - thin borders on all table cells

The DR.md structure it expects (the format produced by the DR-analysis flow):
  <title line>
  Development Review (DR)
  DR ID: ...
  | meta key | meta value |        <- first markdown table (2 cols)
  ### Раздел A — ...
  <intro paragraph(s)>
  | header | ... |                  <- markdown table
  ### Раздел B — ...
  | header | ... |
  ### Раздел C — ...
  <paragraph(s)>
  1. recommendation
  2. ...
Lines starting with '>' become a small italic note.
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE_HDR = "1F4E79"      # Раздел A/B header fill
GREY_LBL = "D9D9D9"      # meta label fill
HEAD_BLUE = RGBColor(0x2F, 0x75, 0xB5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '808080')
        borders.append(el)
    tblPr.append(borders)


def set_cell(cell, text, *, bold=False, color=None, fill=None, size=None):
    cell.text = ""
    p = cell.paragraphs[0]
    parts = text.split("<br>")
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        r = p.add_run(part)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    if fill is not None:
        shade(cell, fill)


def parse_md_table(lines, i):
    """lines[i] starts a md table. Return (rows, next_i). Skips the |---| separator."""
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(set(c) <= set("-: ") for c in cells):  # skip separator row
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows, *, meta=False):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    set_borders(t)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            if meta:
                # 2-col key/value: left column grey+bold label
                if ci == 0:
                    set_cell(cells[ci], txt, bold=True, fill=GREY_LBL)
                else:
                    set_cell(cells[ci], txt)
            else:
                if ri == 0:  # header row dark blue
                    set_cell(cells[ci], txt, bold=True, color=WHITE, fill=BLUE_HDR)
                else:
                    set_cell(cells[ci], txt)
    return t


def render(md_path, docx_path):
    with open(md_path) as f:
        lines = [ln.rstrip("\n") for ln in f]
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)
    i = 0
    title_done = False
    first_table = True
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows, meta=first_table)
            first_table = False
            doc.add_paragraph("")
            continue
        if s.startswith("###"):
            h = doc.add_heading(s.lstrip("#").strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = HEAD_BLUE
                r.font.bold = True
            i += 1
            continue
        if s.startswith(">"):
            p = doc.add_paragraph()
            r = p.add_run(s.lstrip(">").strip())
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            i += 1
            continue
        # plain paragraph
        if not title_done:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.bold = True
            r.font.size = Pt(18)
            title_done = True
        else:
            doc.add_paragraph(s)
        i += 1
    doc.save(docx_path)
    print(f"{md_path} -> {docx_path}")


if __name__ == "__main__":
    render(sys.argv[1], sys.argv[2])
